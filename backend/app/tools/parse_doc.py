"""Parse .docx → text via python-docx."""
from __future__ import annotations

from docx import Document


def parse_docx(path: str) -> dict:
  doc = Document(path)
  paragraphs = []
  for p in doc.paragraphs:
    if p.text.strip():
      paragraphs.append(p.text)
  # also walk tables
  for table in doc.tables:
    for row in table.rows:
      cells = [c.text.strip() for c in row.cells if c.text.strip()]
      if cells:
        paragraphs.append(" | ".join(cells))

  content = "\n\n".join(paragraphs).strip()
  title = ""
  if doc.core_properties and doc.core_properties.title:
    title = str(doc.core_properties.title)[:200]
  if not title:
    # fall back to first non-empty paragraph
    for p in doc.paragraphs:
      if p.text.strip():
        title = p.text.strip()[:80]
        break
  if not title:
    title = "Untitled Document"
  return {"title": title, "content": content or "(empty document)"}